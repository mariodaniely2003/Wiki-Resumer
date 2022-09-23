import wikipedia 
import re 
from docx import Document

nome = input('Como se chama? (o nome será também o autor do documento): ')
wikipedia.set_lang('pt')
titulo = input('Qual é o tema do trabalho? (o tema será também o titulo do documento): ')
while True:
    try:
        wiki = wikipedia.page(titulo)
        break
    except:
        print('Tema inválido ou não encontrado')
        titulo = input('Qual é o tema do trabalho? (o tema será também o titulo do documento): ')

texto = wiki.content
texto = re.sub(r'==','',texto)
texto = re.sub(r'=','',texto)
texto = re.sub(r'\n','\n',texto)
split = texto.split('Veja também',1)
texto = split[0]

print(texto)
document = Document()
paragrafo = document.add_heading(titulo,0)
paragrafo.alignment = 1

paragrafo = document.add_paragraph('    ' + texto)
paragrafo = document.add_paragraph(nome)
paragrafo.alignment = 2
document.save(titulo + ".docx")
input()